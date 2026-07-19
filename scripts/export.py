#!/usr/bin/env python3
"""
导出模块：将 DeepSeek 聊天记录导出为 Markdown、Word、PDF、JSON、JSONL 格式。
支持 YAML Front Matter 元数据、标签系统、全文搜索索引、逐会话 ZIP。
"""

import json
import os
import re
import sqlite3
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional
from html import escape

import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


def load_config(config_path: str = None) -> dict:
    if config_path is None:
        # 查找 config.yaml: 先查 scripts 目录，再查父目录
        script_dir = Path(__file__).parent
        config_path = script_dir / "config.yaml"
        if not config_path.exists():
            config_path = script_dir.parent / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_all_chats(backup_dir: Path) -> List[dict]:
    """加载所有已备份的聊天 JSON 文件。"""
    chats = []
    json_dir = backup_dir / "json"
    if not json_dir.exists():
        return chats
    for f in sorted(json_dir.glob("*.json")):
        try:
            with open(f, "r", encoding="utf-8") as fp:
                chats.append(json.load(fp))
        except Exception:
            pass
    return chats


def filter_chats_by_date(chats: List[dict], from_date: str = None, to_date: str = None) -> List[dict]:
    """按日期范围筛选对话。日期格式: YYYY-MM-DD"""
    if not from_date and not to_date:
        return chats

    filtered = []
    for chat in chats:
        scraped = chat.get("scraped_at", "")
        if not scraped:
            filtered.append(chat)
            continue

        try:
            # 解析 ISO 格式日期
            chat_date = datetime.fromisoformat(scraped.replace("Z", "+00:00")).strftime("%Y-%m-%d")
        except Exception:
            filtered.append(chat)
            continue

        if from_date and chat_date < from_date:
            continue
        if to_date and chat_date > to_date:
            continue
        filtered.append(chat)

    return filtered


def filter_chats_by_keyword(chats: List[dict], keyword: str = None) -> List[dict]:
    """按关键词筛选对话标题。"""
    if not keyword:
        return chats
    keyword_lower = keyword.lower()
    return [c for c in chats if keyword_lower in c.get("title", "").lower()]


# ========== Markdown 导出 ==========

def export_markdown(chats: List[dict], output_path: Path, config: dict):
    lines = [
        "# DeepSeek 聊天记录备份\n",
        f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n",
        f"共 {len(chats)} 个对话\n",
        "---\n",
    ]

    for chat in chats:
        title = chat.get("title", "未命名对话")
        chat_id = chat.get("chat_id", "")
        scraped_at = chat.get("scraped_at", "")
        messages = chat.get("messages", [])

        lines.append(f"\n## {title}\n")
        lines.append(f"- **对话ID**: `{chat_id}`")
        if scraped_at:
            lines.append(f"- **抓取时间**: {scraped_at}")
        lines.append(f"- **消息数**: {len(messages)}\n")

        for msg in messages:
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            role_label = "🤖 Assistant" if role == "assistant" else "👤 User" if role == "user" else f"❓ {role}"
            lines.append(f"### {role_label}\n")
            lines.append(f"{content}\n")
            lines.append("---\n")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[✓] Markdown: {output_path}")


# ========== Word 导出 ==========

def export_word(chats: List[dict], output_path: Path, config: dict):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = config.get("word", {}).get("font_name", "Arial")
    style.font.size = Pt(config.get("word", {}).get("font_size", 11))

    title = doc.add_heading("DeepSeek 聊天记录备份", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 共 {len(chats)} 个对话")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # 目录
    doc.add_heading("目录", level=1)
    for i, chat in enumerate(chats, 1):
        doc.add_paragraph(f"{i}. {chat.get('title', '未命名对话')}")
    doc.add_page_break()

    # 内容
    for chat in chats:
        doc.add_heading(chat.get("title", "未命名对话"), level=1)
        meta_para = doc.add_paragraph()
        if chat.get("chat_id"):
            run = meta_para.add_run(f"ID: {chat['chat_id']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
        if chat.get("scraped_at"):
            run = meta_para.add_run(f"  |  {chat['scraped_at']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        for msg in chat.get("messages", []):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            if role == "assistant":
                role_label, color = "Assistant", RGBColor(0, 100, 200)
            elif role == "user":
                role_label, color = "User", RGBColor(0, 150, 0)
            else:
                role_label, color = role, RGBColor(128, 128, 128)

            h = doc.add_heading(level=2)
            run = h.add_run(role_label)
            run.font.color.rgb = color
            run.font.size = Pt(12)
            p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(12)
            doc.add_paragraph("─" * 40)

        doc.add_page_break()

    doc.save(str(output_path))
    print(f"[✓] Word: {output_path}")


# ========== PDF 导出 ==========

if HAS_PDF:
    class PDFExporter(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

        def header(self):
            self.set_font("Helvetica", "I", 8)
            self.cell(0, 10, "DeepSeek Chat Backup", align="R")
            self.ln(10)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        def add_title_page(self, count: int):
            self.add_page()
            self.set_font("Helvetica", "B", 24)
            self.cell(0, 40, "", ln=True)
            self.cell(0, 15, "DeepSeek Chat Backup", align="C", ln=True)
            self.ln(10)
            self.set_font("Helvetica", "", 12)
            self.cell(0, 10, f"Export Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align="C", ln=True)
            self.cell(0, 10, f"Total Conversations: {count}", align="C", ln=True)

        def add_chat(self, chat: dict):
            self.add_page()
            self.set_font("Helvetica", "B", 16)
            title = chat.get("title", "Untitled")[:80]
            self.multi_cell(0, 10, title)
            self.ln(5)

            self.set_font("Helvetica", "I", 9)
            if chat.get("scraped_at"):
                self.cell(0, 6, f"Scraped: {chat['scraped_at']}", ln=True)
            self.cell(0, 6, f"Messages: {len(chat.get('messages', []))}", ln=True)
            self.ln(5)

            for msg in chat.get("messages", []):
                role = msg.get("role", "unknown")
                content = msg.get("content", "")

                if role == "assistant":
                    self.set_fill_color(230, 240, 255)
                    role_label = "Assistant"
                elif role == "user":
                    self.set_fill_color(230, 255, 230)
                    role_label = "User"
                else:
                    self.set_fill_color(245, 245, 245)
                    role_label = role

                self.set_font("Helvetica", "B", 10)
                self.cell(0, 8, f"  {role_label}", fill=True, ln=True)
                self.ln(2)

                self.set_font("Helvetica", "", 10)
                safe_content = content.encode("latin-1", errors="replace").decode("latin-1")
                self.multi_cell(0, 6, safe_content)
                self.ln(5)

                self.set_draw_color(200, 200, 200)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(5)


def export_pdf(chats: List[dict], output_path: Path, config: dict):
    if not HAS_PDF:
        print("[!] PDF 导出需要 fpdf2: pip install fpdf2")
        return
    pdf = PDFExporter()
    pdf.alias_nb_pages()
    pdf.add_title_page(len(chats))

    for chat in chats:
        pdf.add_chat(chat)

    pdf.output(str(output_path))
    print(f"[✓] PDF: {output_path}")


# ========== JSON 导出 ==========

def export_json(chats: List[dict], output_path: Path, config: dict):
    data = {
        "export_time": datetime.now(timezone.utc).isoformat(),
        "count": len(chats),
        "conversations": chats,
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"[✓] JSON: {output_path}")


# ========== HTML Viewer 导出 ==========

def export_html_viewer(chats: List[dict], output_dir: Path, config: dict):
    """生成可浏览的 HTML 页面（借鉴 chatgpt-exporter 的设计）。"""
    html_dir = output_dir / "html_viewer"
    html_dir.mkdir(exist_ok=True)

    # 生成主页面
    stats = get_chat_stats(chats)
    sidebar_items = ""
    for i, chat in enumerate(chats):
        title = escape(chat.get("title", "Untitled")[:60])
        msg_count = len(chat.get("messages", []))
        safe_id = f"chat_{i}"
        sidebar_items += f'<div class="sidebar-item" onclick="showChat(\'{safe_id}\')">{title} <span class="msg-count">({msg_count})</span></div>\n'

    # 生成每个对话的内容
    chat_contents = ""
    for i, chat in enumerate(chats):
        title = escape(chat.get("title", "Untitled"))
        messages_html = ""
        for j, msg in enumerate(chat.get("messages", [])):
            role = msg.get("role", "unknown")
            content = escape(msg.get("content", ""))
            role_class = "assistant" if role == "assistant" else "user"
            role_label = "🤖 Assistant" if role == "assistant" else "👤 User"
            
            # 长消息折叠：超过 200 字符的消息默认折叠
            is_long = len(content) > 200
            content_class = "content collapsed" if is_long else "content"
            expand_btn = f'<button class="expand-btn" onclick="toggleMessage(this)">展开全文</button>' if is_long else ''
            
            messages_html += f'''<div class="message {role_class}">
                <div class="role">{role_label}</div>
                <div class="{content_class}">{content}</div>
                {expand_btn}
            </div>\n'''
        chat_contents += f'<div class="chat-content" id="chat_{i}" style="display:none"><h2>{title}</h2>{messages_html}</div>\n'

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>DeepSeek Chat Backup</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; display: flex; height: 100vh; background: #1a1a2e; color: #eee; }}
.sidebar {{ width: 300px; background: #16213e; overflow-y: auto; border-right: 1px solid #0f3460; }}
.sidebar-header {{ padding: 20px; background: #0f3460; text-align: center; }}
.sidebar-header h2 {{ font-size: 16px; margin-bottom: 10px; }}
.sidebar-stats {{ font-size: 12px; color: #aaa; }}
.sidebar-item {{ padding: 12px 16px; cursor: pointer; border-bottom: 1px solid #1a1a3e; font-size: 14px; transition: background 0.2s; }}
.sidebar-item:hover {{ background: #1a1a3e; }}
.sidebar-item.active {{ background: #0f3460; }}
.msg-count {{ color: #888; font-size: 12px; }}
.main {{ flex: 1; overflow-y: auto; padding: 20px; }}
.chat-content h2 {{ margin-bottom: 20px; color: #e94560; }}
.message {{ margin-bottom: 16px; padding: 12px 16px; border-radius: 8px; max-width: 80%; }}
.message.user {{ background: #1a3a5c; margin-right: auto; }}
.message.assistant {{ background: #2a1a3e; margin-left: auto; }}
.role {{ font-size: 12px; color: #aaa; margin-bottom: 4px; }}
.content {{ font-size: 14px; line-height: 1.6; white-space: pre-wrap; word-wrap: break-word; }}
.content.collapsed {{ max-height: 120px; overflow: hidden; position: relative; }}
.content.collapsed::after {{ content: ""; position: absolute; bottom: 0; left: 0; right: 0; height: 40px; background: linear-gradient(transparent, #2a1a3e); pointer-events: none; }}
.message.user .content.collapsed::after {{ background: linear-gradient(transparent, #1a3a5c); }}
.expand-btn {{ display: inline-block; margin-top: 8px; padding: 4px 12px; font-size: 12px; color: #e94560; background: transparent; border: 1px solid #e94560; border-radius: 4px; cursor: pointer; transition: all 0.2s; }}
.expand-btn:hover {{ background: #e94560; color: #fff; }}
.search {{ padding: 12px; background: #16213e; }}
.search input {{ width: 100%; padding: 8px 12px; border: 1px solid #0f3460; border-radius: 4px; background: #1a1a2e; color: #eee; }}
</style>
</head>
<body>
<div class="sidebar">
<div class="sidebar-header">
<h2>DeepSeek Chat Backup</h2>
<div class="sidebar-stats">共 {stats['total_chats']} 个对话 | {stats['total_messages']} 条消息 | {stats['total_words']:,} 字</div>
</div>
<div class="search"><input type="text" placeholder="搜索对话..." oninput="filterSidebar(this.value)"></div>
<div class="sidebar-list">{sidebar_items}</div>
</div>
<div class="main">
{chat_contents}
<div id="welcome" style="text-align:center;margin-top:100px;color:#666;">
<h2>← 选择一个对话开始浏览</h2>
</div>
</div>
<script>
function showChat(id) {{
    document.querySelectorAll('.chat-content').forEach(el => el.style.display = 'none');
    document.getElementById(id).style.display = 'block';
    document.getElementById('welcome').style.display = 'none';
    document.querySelectorAll('.sidebar-item').forEach(el => el.classList.remove('active'));
    event.currentTarget.classList.add('active');
}}
function filterSidebar(query) {{
    const items = document.querySelectorAll('.sidebar-item');
    items.forEach(item => {{
        item.style.display = item.textContent.toLowerCase().includes(query.toLowerCase()) ? 'block' : 'none';
    }});
}}
function toggleMessage(btn) {{
    const content = btn.previousElementSibling;
    if (content.classList.contains('collapsed')) {{
        content.classList.remove('collapsed');
        btn.textContent = '收起';
    }} else {{
        content.classList.add('collapsed');
        btn.textContent = '展开全文';
    }}
}}
</script>
</body>
</html>'''

    with open(html_dir / "index.html", "w", encoding="utf-8") as f:
        f.write(html)
    print(f"[✓] HTML Viewer: {html_dir / 'index.html'}")


# ========== 对话统计 ==========

def get_chat_stats(chats: List[dict]) -> dict:
    """计算对话统计信息。"""
    total_messages = sum(len(c.get("messages", [])) for c in chats)
    total_words = sum(
        sum(len(m.get("content", "")) for m in c.get("messages", []))
        for c in chats
    )
    user_words = sum(
        sum(len(m.get("content", "")) for m in c.get("messages", []) if m.get("role") == "user")
        for c in chats
    )
    assistant_words = total_words - user_words
    return {
        "total_chats": len(chats),
        "total_messages": total_messages,
        "total_words": total_words,
        "user_words": user_words,
        "assistant_words": assistant_words,
    }


# ========== ZIP 归档 ==========

def create_zip_archive(export_dir: Path, timestamp: str) -> Path:
    """将导出目录打包为 ZIP。"""
    zip_path = export_dir / f"deepseek_backup_{timestamp}.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in export_dir.iterdir():
            if f.is_file() and f.suffix in ['.md', '.docx', '.pdf', '.json']:
                zf.write(f, f.name)
            elif f.is_dir() and f.name == 'html_viewer':
                for hf in f.rglob('*'):
                    if hf.is_file():
                        zf.write(hf, f"html_viewer/{hf.name}")
    print(f"[✓] ZIP: {zip_path}")
    return zip_path


# ========== JSONL 导出（训练数据格式）==========

def export_jsonl(chats: List[dict], output_path: Path, config: dict):
    """导出为 JSONL 格式（每行一个 JSON，适合 AI 训练）。"""
    with open(output_path, "w", encoding="utf-8") as f:
        for chat in chats:
            title = chat.get("title", "")
            chat_id = chat.get("chat_id", "")
            scraped_at = chat.get("scraped_at", "")
            messages = chat.get("messages", [])
            
            for msg in messages:
                record = {
                    "chat_id": chat_id,
                    "chat_title": title,
                    "timestamp": scraped_at,
                    "role": msg.get("role", "unknown"),
                    "content": msg.get("content", ""),
                }
                f.write(json.dumps(record, ensure_ascii=False) + "\n")
    print(f"[✓] JSONL: {output_path}")


# ========== YAML Front Matter Markdown 导出 ==========

def export_markdown_with_frontmatter(chats: List[dict], output_path: Path, config: dict):
    """导出带 YAML Front Matter 的 Markdown（借鉴 Claude Extractor）。"""
    stats = get_chat_stats(chats)
    lines = [
        "---",
        f"title: DeepSeek Chat Backup",
        f"date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"conversations: {stats['total_chats']}",
        f"messages: {stats['total_messages']}",
        f"words: {stats['total_words']}",
        f"tags: [deepseek, backup, chat-history]",
        "---\n",
        "# DeepSeek 聊天记录备份\n",
    ]

    for chat in chats:
        title = chat.get("title", "未命名对话")
        chat_id = chat.get("chat_id", "")
        scraped_at = chat.get("scraped_at", "")
        messages = chat.get("messages", [])
        tags = chat.get("tags", [])

        lines.append(f"\n## {title}\n")
        lines.append(f"```yaml")
        lines.append(f"id: {chat_id}")
        if scraped_at:
            lines.append(f"date: {scraped_at[:10]}")
        lines.append(f"messages: {len(messages)}")
        if tags:
            lines.append(f"tags: [{', '.join(tags)}]")
        lines.append(f"```\n")

        for msg in messages:
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            role_label = "🤖 Assistant" if role == "assistant" else "👤 User" if role == "user" else f"❓ {role}"
            lines.append(f"### {role_label}\n")
            lines.append(f"{content}\n")
            lines.append("---\n")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[✓] Markdown (Front Matter): {output_path}")


# ========== SQLite 全文搜索索引 ==========

def build_search_index(chats: List[dict], db_path: Path):
    """构建 SQLite FTS5 全文搜索索引（借鉴 Kept 项目）。"""
    if db_path.exists():
        db_path.unlink()
    
    conn = sqlite3.connect(str(db_path))
    cursor = conn.cursor()
    
    # 创建 FTS5 虚拟表
    cursor.execute("""
        CREATE VIRTUAL TABLE IF NOT EXISTS chat_fts 
        USING fts5(
            chat_id, 
            title, 
            content, 
            role,
            tokenize='unicode61'
        )
    """)
    
    # 创建普通表存储元数据
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS chat_meta (
            chat_id TEXT PRIMARY KEY,
            title TEXT,
            scraped_at TEXT,
            message_count INTEGER,
            content_hash TEXT,
            tags TEXT
        )
    """)
    
    # 插入数据
    for chat in chats:
        chat_id = chat.get("chat_id", "")
        title = chat.get("title", "")
        scraped_at = chat.get("scraped_at", "")
        messages = chat.get("messages", [])
        content_hash = chat.get("content_hash", "")
        tags = chat.get("tags", [])
        
        # 插入元数据
        cursor.execute("""
            INSERT OR REPLACE INTO chat_meta 
            (chat_id, title, scraped_at, message_count, content_hash, tags)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (chat_id, title, scraped_at, len(messages), content_hash, json.dumps(tags)))
        
        # 插入 FTS 索引
        for msg in messages:
            cursor.execute("""
                INSERT INTO chat_fts (chat_id, title, content, role)
                VALUES (?, ?, ?, ?)
            """, (chat_id, title, msg.get("content", ""), msg.get("role", "unknown")))
    
    conn.commit()
    conn.close()
    print(f"[✓] Search Index: {db_path}")


def search_chats(db_path: Path, query: str, limit: int = 20) -> List[dict]:
    """全文搜索对话内容。"""
    if not db_path.exists():
        return []
    
    conn = sqlite3.connect(str(db_path))
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT DISTINCT chat_id, title, 
               snippet(chat_fts, 2, '<mark>', '</mark>', '...', 32) as snippet
        FROM chat_fts 
        WHERE chat_fts MATCH ? 
        LIMIT ?
    """, (query, limit))
    
    results = []
    for row in cursor.fetchall():
        results.append({
            "chat_id": row[0],
            "title": row[1],
            "snippet": row[2],
        })
    
    conn.close()
    return results


# ========== 逐会话 ZIP（借鉴 DeepSeek Exporter）==========

def create_per_chat_zips(chats: List[dict], output_dir: Path):
    """为每个对话创建独立的 ZIP 文件。"""
    per_chat_dir = output_dir / "per_chat_zips"
    per_chat_dir.mkdir(exist_ok=True)
    
    for chat in chats:
        chat_id = chat.get("chat_id", "")
        title = chat.get("title", "untitled")
        safe_title = "".join(c if c.isalnum() or c in "-_" else "_" for c in title)[:50]
        safe_id = "".join(c if c.isalnum() or c in "-_" else "_" for c in chat_id)[:8]
        
        zip_path = per_chat_dir / f"{safe_title}_{safe_id}.zip"
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            # 写入 JSON
            zf.writestr("chat.json", json.dumps(chat, ensure_ascii=False, indent=2))
            
            # 写入 Markdown
            md_content = f"# {title}\n\n"
            for msg in chat.get("messages", []):
                role = msg.get("role", "unknown")
                content = msg.get("content", "")
                role_label = "Assistant" if role == "assistant" else "User"
                md_content += f"### {role_label}\n\n{content}\n\n---\n\n"
            zf.writestr("chat.md", md_content)
    
    print(f"[✓] Per-chat ZIPs: {per_chat_dir} ({len(chats)} files)")


# ========== 标签管理 ==========

def load_tags(config: dict) -> dict:
    """加载标签配置。"""
    backup_dir = Path(os.path.expanduser(config["backup_dir"]))
    tags_file = backup_dir / "tags.json"
    if tags_file.exists():
        with open(tags_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_tags(config: dict, tags: dict):
    """保存标签配置。"""
    backup_dir = Path(os.path.expanduser(config["backup_dir"]))
    tags_file = backup_dir / "tags.json"
    with open(tags_file, "w", encoding="utf-8") as f:
        json.dump(tags, f, ensure_ascii=False, indent=2)


def add_tag_to_chat(config: dict, chat_id: str, tag: str):
    """为对话添加标签。"""
    tags = load_tags(config)
    if chat_id not in tags:
        tags[chat_id] = []
    if tag not in tags[chat_id]:
        tags[chat_id].append(tag)
    save_tags(config, tags)


def filter_chats_by_tag(chats: List[dict], tag: str) -> List[dict]:
    """按标签筛选对话。"""
    if not tag:
        return chats
    # 从对话数据中筛选
    return [c for c in chats if tag in c.get("tags", [])]


# ========== 统一导出入口 ==========

FORMAT_HANDLERS = {
    "markdown": export_markdown,
    "md": export_markdown,
    "word": export_word,
    "docx": export_word,
    "pdf": export_pdf,
    "json": export_json,
    "jsonl": export_jsonl,
    "html": export_html_viewer,
    "frontmatter": export_markdown_with_frontmatter,
}

FORMAT_EXTENSIONS = {
    "markdown": ".md",
    "md": ".md",
    "word": ".docx",
    "docx": ".docx",
    "pdf": ".pdf",
    "json": ".json",
    "jsonl": ".jsonl",
    "html": ".html",
    "frontmatter": ".md",
}


def export_all(
    config: dict = None,
    formats: List[str] = None,
    from_date: str = None,
    to_date: str = None,
    keyword: str = None,
    chat_ids: List[str] = None,
    tag: str = None,
    create_zip: bool = False,
    per_chat_zip: bool = False,
    build_index: bool = False,
):
    """统一导出入口。"""
    if config is None:
        config = load_config()
    if formats is None:
        formats = config.get("export_formats", ["markdown", "word"])

    backup_dir = Path(os.path.expanduser(config["backup_dir"]))
    chats = load_all_chats(backup_dir)

    if not chats:
        print("[✗] 没有找到已备份的聊天记录")
        return

    # 筛选
    if from_date or to_date:
        chats = filter_chats_by_date(chats, from_date, to_date)
        print(f"[i] 日期筛选后: {len(chats)} 个对话")

    if keyword:
        chats = filter_chats_by_keyword(chats, keyword)
        print(f"[i] 关键词筛选后: {len(chats)} 个对话")

    if chat_ids:
        id_set = set(chat_ids)
        chats = [c for c in chats if c.get("chat_id") in id_set]
        print(f"[i] ID 筛选后: {len(chats)} 个对话")

    if tag:
        chats = filter_chats_by_tag(chats, tag)
        print(f"[i] 标签筛选后: {len(chats)} 个对话")

    if not chats:
        print("[✗] 筛选后无匹配对话")
        return

    # 统计信息
    stats = get_chat_stats(chats)
    print(f"[i] 导出 {stats['total_chats']} 个对话 ({stats['total_messages']} 条消息, {stats['total_words']:,} 字)")

    export_dir = backup_dir / "exports"
    export_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    for fmt in formats:
        fmt_lower = fmt.lower()
        handler = FORMAT_HANDLERS.get(fmt_lower)
        ext = FORMAT_EXTENSIONS.get(fmt_lower, ".bin")
        if handler:
            if fmt_lower in ["html"]:
                handler(chats, export_dir, config)
            else:
                output_path = export_dir / f"deepseek_backup_{timestamp}{ext}"
                handler(chats, output_path, config)
        else:
            print(f"[!] 不支持的格式: {fmt}")

    # 构建搜索索引
    if build_index:
        db_path = export_dir / f"deepseek_search_{timestamp}.db"
        build_search_index(chats, db_path)

    # 逐会话 ZIP
    if per_chat_zip:
        create_per_chat_zips(chats, export_dir)

    # ZIP 归档
    if create_zip:
        create_zip_archive(export_dir, timestamp)

    print(f"\n[✓] 导出完成！目录: {export_dir}")


def main():
    import argparse
    parser = argparse.ArgumentParser(description="DeepSeek 聊天记录导出工具")
    parser.add_argument("--format", "-f", nargs="+",
                        choices=["markdown", "md", "word", "docx", "pdf", "json", "jsonl", "html", "frontmatter", "all"],
                        default=["all"], help="导出格式 (可多选)")
    parser.add_argument("--from-date", help="起始日期 YYYY-MM-DD")
    parser.add_argument("--to-date", help="截止日期 YYYY-MM-DD")
    parser.add_argument("--keyword", "-k", help="按标题关键词筛选")
    parser.add_argument("--tag", "-t", help="按标签筛选")
    parser.add_argument("--chat-id", nargs="+", help="指定对话 ID (可多选)")
    parser.add_argument("--list", "-l", action="store_true", help="列出所有对话")
    parser.add_argument("--stats", "-s", action="store_true", help="显示统计信息")
    parser.add_argument("--search", help="全文搜索对话内容")
    parser.add_argument("--zip", "-z", action="store_true", help="打包为 ZIP")
    parser.add_argument("--per-chat-zip", action="store_true", help="逐会话 ZIP")
    parser.add_argument("--build-index", action="store_true", help="构建搜索索引")
    parser.add_argument("--config", "-c", help="配置文件路径")
    args = parser.parse_args()

    config = load_config(args.config)

    # 全文搜索
    if args.search:
        backup_dir = Path(os.path.expanduser(config["backup_dir"]))
        db_files = list(backup_dir.glob("exports/*.db"))
        if not db_files:
            print("[✗] 未找到搜索索引，请先运行: export.py --build-index")
            return
        db_path = max(db_files, key=os.path.getmtime)
        results = search_chats(db_path, args.search)
        print(f"\n搜索 '{args.search}' 结果 ({len(results)} 条):")
        for r in results:
            print(f"  [{r['title']}] {r['snippet'][:80]}...")
        return

    # 列出对话
    if args.list:
        backup_dir = Path(os.path.expanduser(config["backup_dir"]))
        chats = load_all_chats(backup_dir)
        if from_date := args.from_date:
            chats = filter_chats_by_date(chats, from_date=from_date)
        if to_date := args.to_date:
            chats = filter_chats_by_date(chats, to_date=to_date)
        if keyword := args.keyword:
            chats = filter_chats_by_keyword(chats, keyword)
        if tag := args.tag:
            chats = filter_chats_by_tag(chats, tag)

        print(f"\n{'ID':<45} {'日期':<12} {'标题'}")
        print("-" * 100)
        for c in chats:
            cid = c.get("chat_id", "")[:43]
            date = c.get("scraped_at", "")[:10]
            title = c.get("title", "")[:40]
            print(f"{cid:<45} {date:<12} {title}")
        print(f"\n共 {len(chats)} 个对话")
        return

    # 统计信息
    if args.stats:
        backup_dir = Path(os.path.expanduser(config["backup_dir"]))
        chats = load_all_chats(backup_dir)
        if from_date := args.from_date:
            chats = filter_chats_by_date(chats, from_date=from_date)
        if to_date := args.to_date:
            chats = filter_chats_by_date(chats, to_date=to_date)
        if keyword := args.keyword:
            chats = filter_chats_by_keyword(chats, keyword)

        stats = get_chat_stats(chats)
        print(f"\n{'='*50}")
        print(f"DeepSeek 聊天记录统计")
        print(f"{'='*50}")
        print(f"  对话总数: {stats['total_chats']}")
        print(f"  消息总数: {stats['total_messages']}")
        print(f"  总字数:   {stats['total_words']:,}")
        print(f"  用户字数: {stats['user_words']:,}")
        print(f"  AI字数:   {stats['assistant_words']:,}")
        print(f"{'='*50}")
        return

    # 处理格式
    formats = []
    for f in args.format:
        if f == "all":
            formats = ["markdown", "word", "json", "jsonl", "html", "frontmatter"]
            break
        formats.append(f)

    export_all(
        config=config,
        formats=formats,
        from_date=args.from_date,
        to_date=args.to_date,
        keyword=args.keyword,
        chat_ids=args.chat_id,
        tag=args.tag,
        create_zip=args.zip,
        per_chat_zip=args.per_chat_zip,
        build_index=args.build_index,
    )


if __name__ == "__main__":
    main()
